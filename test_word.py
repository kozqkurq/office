from docx import Document

document = Document()

document.add_heading('簡単なWordドキュメントのタイトル', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

document.save('sample.docx')


# 演習
doc = Document("sample.docx")

doc.add_picture("./anya.jpg")

count = 0
for para in doc.paragraphs:
    count += len(para.text)

doc.add_paragraph(f"文字数は{count}個")

doc.save('sample_answer.docx')